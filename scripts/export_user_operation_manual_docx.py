# ============================================================
# File Name   : export_user_operation_manual_docx.py
# Description:
#   生成带真实页面截图的“数语”用户操作手册 DOCX。
#
# Responsibilities:
#   - 读取 docs/user-manual-screenshots 下的真实页面截图并嵌入 Word 文档。
#   - 按功能页组织“功能用途、入口位置、操作步骤、结果确认、注意事项”，形成面向最终用户的操作手册。
#
# Author      : yangkai
# Created On  : 2026-06-12
# ============================================================

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.export_user_manual_docx import (
    BLUE,
    BLUE_GRAY_FILL,
    DARK_BLUE,
    INK,
    LIGHT_FILL,
    MUTED,
    add_bullets,
    add_callout,
    add_heading,
    add_para,
    add_steps,
    configure_styles,
    set_header_footer,
    set_run_font,
    style_table,
)


OUTPUT_PATH = ROOT / "docs" / "数语多智能体产品用户手册.docx"
SCREENSHOT_DIR = ROOT / "docs" / "user-manual-screenshots"


SCREENS = {
    "workspace": SCREENSHOT_DIR / "01-workspace.png",
    "chat": SCREENSHOT_DIR / "02-chat-empty.png",
    "datasets": SCREENSHOT_DIR / "03-datasets-blueprints.png",
    "datasources": SCREENSHOT_DIR / "04-datasources.png",
    "audit": SCREENSHOT_DIR / "05-audit-query.png",
    "settings": SCREENSHOT_DIR / "06-settings.png",
    "history": SCREENSHOT_DIR / "07-history.png",
    "apis": SCREENSHOT_DIR / "08-apis.png",
    "dataset_blueprints_tab": SCREENSHOT_DIR / "09-dataset-analysis-blueprints-tab.png",
    "dataset_tables_tab": SCREENSHOT_DIR / "10-dataset-tables-tab.png",
    "dataset_field_annotation_tab": SCREENSHOT_DIR / "11-dataset-field-annotation-tab.png",
    "dataset_metrics_tab": SCREENSHOT_DIR / "12-dataset-metrics-tab.png",
    "dataset_dimensions_tab": SCREENSHOT_DIR / "13-dataset-dimensions-tab.png",
    "dataset_validation_tab": SCREENSHOT_DIR / "14-dataset-validation-tab.png",
    "dataset_manifest_tab": SCREENSHOT_DIR / "15-dataset-manifest-tab.png",
    "dataset_terms_tab": SCREENSHOT_DIR / "16-dataset-terms-tab.png",
    "dataset_scenarios_tab": SCREENSHOT_DIR / "17-dataset-scenarios-tab.png",
    "dataset_permissions_tab": SCREENSHOT_DIR / "18-dataset-permissions-tab.png",
    "dataset_version_history_tab": SCREENSHOT_DIR / "19-dataset-version-history-tab.png",
    "datasource_overview_tab": SCREENSHOT_DIR / "20-datasource-overview-tab.png",
    "datasource_schema_tab": SCREENSHOT_DIR / "21-datasource-schema-tab.png",
    "datasource_ddl_sync_tab": SCREENSHOT_DIR / "22-datasource-ddl-sync-tab.png",
    "api_overview_tab": SCREENSHOT_DIR / "23-api-overview-tab.png",
    "api_params_tab": SCREENSHOT_DIR / "24-api-params-tab.png",
    "api_examples_tab": SCREENSHOT_DIR / "25-api-examples-tab.png",
    "api_logs_tab": SCREENSHOT_DIR / "26-api-logs-tab.png",
    "api_versions_tab": SCREENSHOT_DIR / "27-api-versions-tab.png",
    "settings_profile": SCREENSHOT_DIR / "28-settings-profile.png",
    "settings_notifications": SCREENSHOT_DIR / "29-settings-notifications.png",
    "settings_appearance": SCREENSHOT_DIR / "30-settings-appearance.png",
    "settings_workspace": SCREENSHOT_DIR / "31-settings-workspace.png",
    "settings_members_permissions": SCREENSHOT_DIR / "32-settings-members-permissions.png",
    "settings_usage_billing": SCREENSHOT_DIR / "33-settings-usage-billing.png",
    "settings_data_sources": SCREENSHOT_DIR / "34-settings-data-sources.png",
    "settings_llm_models": SCREENSHOT_DIR / "35-settings-llm-models.png",
    "settings_business_dictionary": SCREENSHOT_DIR / "36-settings-business-dictionary.png",
    "settings_api_keys": SCREENSHOT_DIR / "37-settings-api-keys.png",
    "settings_webhooks": SCREENSHOT_DIR / "38-settings-webhooks.png",
    "settings_audit_logs": SCREENSHOT_DIR / "39-settings-audit-logs.png",
    "blueprint_detail_overview_tab": SCREENSHOT_DIR / "40-blueprint-detail-overview-tab.png",
}


def add_cover(doc: Document) -> None:
    add_para(doc, "操作手册", size=11, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=18)
    add_para(doc, "数语多智能体产品用户手册", size=28, bold=True, color=DARK_BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "基于真实页面截图的功能说明与操作步骤", size=14,
             color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)

    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("适用对象", "问数用户、数据分析师、数据管理员、交付培训人员"),
        ("截图来源", "本地运行页面 http://localhost:5173，截图时间 2026-06-12"),
        ("手册目标", "帮助用户知道每个页面有哪些功能、从哪里进入、如何操作、如何判断结果"),
        ("文档格式", "Microsoft Word DOCX"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
    style_table(table, [1800, 7560], header=False)
    for row in table.rows:
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=10, bold=True, color=DARK_BLUE)
    add_para(doc, "", after=50)
    add_para(doc, "维护人：yangkai / KenYang", size=10.5, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "生成日期：2026-06-12", size=10.5, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    items = [
        "1. 使用前先了解：页面结构与角色分工",
        "2. 工作台：进入系统后的总览页",
        "3. 对话问数：提出问题并查看 AI 回答",
        "4. 数据集 & 指标：维护语义能力和分析蓝图",
        "5. 数据源：接入数据库并同步表结构",
        "6. 查询审计：回看 Trace、SQL、成本和失败原因",
        "7. 查询历史：复用过往会话",
        "8. API 接口：把问数能力发布给系统调用",
        "9. 系统设置：账号、安全、模型与集成配置",
        "10. 常见操作检查清单",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)
    doc.add_page_break()


def add_screenshot(doc: Document, key: str, caption: str) -> None:
    path = SCREENS[key]
    if not path.exists():
        add_callout(doc, "截图缺失", f"未找到截图文件：{path}", fill=LIGHT_FILL)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(caption)
    set_run_font(run, size=9.5, bold=True, color=MUTED)
    doc.add_picture(str(path), width=Inches(6.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "", after=4)


def add_screenshot_series(doc: Document, screenshots: list[tuple[str, str]]) -> None:
    for key, caption in screenshots:
        add_screenshot(doc, key, caption)


def add_function_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    headers = ["功能", "在哪里", "什么时候用"]
    for idx, text in enumerate(headers):
        table.cell(0, idx).text = text
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text
    style_table(table, [1900, 3200, 4260])


def add_detail_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple[str, ...]],
    widths: list[int],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, text in enumerate(headers):
        table.cell(0, idx).text = text
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text
    style_table(table, widths)
    add_para(doc, "", after=2)


def add_screen_section(
    doc: Document,
    title: str,
    screenshot_key: str,
    caption: str,
    purpose: str,
    steps: list[str],
    confirm: list[str],
    notes: list[str] | None = None,
) -> None:
    add_heading(doc, title, 1)
    add_para(doc, purpose)
    add_screenshot(doc, screenshot_key, caption)
    add_heading(doc, "怎么用", 2)
    add_steps(doc, steps)
    add_heading(doc, "操作完成后看什么", 2)
    add_bullets(doc, confirm)
    if notes:
        add_heading(doc, "注意事项", 2)
        add_bullets(doc, notes)
    doc.add_page_break()


def build_manual() -> Document:
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc.sections[0])
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "1. 使用前先了解：页面结构与角色分工", 1)
    add_para(doc, "数语的日常操作围绕左侧导航展开。普通用户主要使用“对话问数、查询历史、我的收藏”；数据分析师主要使用“数据集 & 指标、知识库、审核队列”；管理员和运维人员主要使用“数据源、查询审计、系统设置”。")
    add_callout(
        doc,
        "本手册的阅读方式",
        "每一章都按真实页面截图说明：先看页面上有哪些区域，再按步骤操作，最后用“操作完成后看什么”判断是否成功。截图来自当前本地运行页面，真实数据、数量和 trace 会随环境变化。",
        BLUE_GRAY_FILL,
    )
    add_function_table(doc, [
        ("新的问数", "左侧顶部按钮", "开始一个新的自然语言问数会话。"),
        ("对话问数", "问数中心", "输入问题、查看回答、SQL、图表和推理过程。"),
        ("数据集 & 指标", "语义治理", "维护数据表、字段、指标、维度、蓝图和 Manifest。"),
        ("数据源", "数据连接", "配置数据库连接、测试连接、同步 Schema。"),
        ("查询审计", "系统管理", "定位失败、回看 trace、检查 SQL 和成本。"),
        ("系统设置", "系统管理", "维护账号、权限、模型、API 密钥和 Webhooks。"),
    ])
    doc.add_page_break()

    add_screen_section(
        doc,
        "2. 工作台：进入系统后的总览页",
        "workspace",
        "截图 1：工作台首页",
        "工作台是进入系统后的总览页，用于快速理解当前工作空间、近期指标、待处理事项和常用入口。",
        [
            "登录后默认进入“工作台”，先查看左侧导航确认当前账号和工作区。",
            "查看页面中的关键指标、待办提醒和快捷入口，判断今天需要处理的问数、治理或审计事项。",
            "如果要开始问数，点击左侧“新的问数”或“对话问数”。",
            "如果要治理语义资产，点击“数据集 & 指标”；如果要排查链路，点击“查询审计”。",
        ],
        [
            "左侧导航高亮当前所在页面。",
            "工作台中的卡片能反映当前系统状态，入口按钮能跳转到对应页面。",
            "右上角通知角标提示还有未处理消息。",
        ],
        ["工作台适合看全局状态，不适合做具体数据治理；具体配置需要进入对应功能页。"],
    )

    add_heading(doc, "3. 对话问数：提出问题并查看 AI 回答", 1)
    add_para(doc, "对话问数是终端用户的核心页面。用户用自然语言描述业务问题，系统会完成意图识别、数据集路由、SQL 生成、SQL 执行、结果解释和可视化建议。")
    add_screenshot(doc, "chat", "截图 2：对话问数首页")
    add_heading(doc, "3.1 页面区域与功能", 2)
    add_detail_table(doc, ["区域", "具体功能", "操作说明"], [
        ("左侧最近对话", "查看和继续历史会话", "点击会话标题可回到原对话；业务主题变化较大时建议点击“新的问数”重新开始。"),
        ("中间欢迎区", "展示可直接点击的示例问题", "适合新用户快速体验问数能力，也可作为问题写法参考。"),
        ("输入框", "输入自然语言问题", "问题尽量包含对象、时间、指标和拆分维度，例如“近 30 天各部门任务完成趋势”。"),
        ("输入区控件", "限定数据集、时间范围和分析深度", "当问题涉及特定业务域时先选择数据集；需要更稳结果时提高分析深度。"),
        ("推理过程", "查看 Agent 执行链路", "用于检查系统是否选对数据集、是否正确生成 DSL/SQL、是否有 SQL Guard 或失败回退。"),
        ("保存到看板", "把稳定回答沉淀为看板素材", "适合周期性查看的固定问题；保存前应确认 SQL 和结果口径。"),
        ("发布为接口", "把稳定问数发布给外部系统调用", "只建议对口径稳定、参数清晰、权限明确的问题使用。"),
    ], [1500, 2600, 5260])
    add_heading(doc, "3.2 标准问数流程", 2)
    add_steps(doc, [
        "点击左侧“新的问数”或进入“对话问数”。",
        "输入业务问题，避免只写“分析一下”“看下情况”这类缺少对象和指标的问题。",
        "按需选择数据集范围、时间范围或分析深度。",
        "点击发送，等待 AI 返回回答、表格、图表、SQL 或风险提示。",
        "点击“推理过程”核对入口识别、数据集路由、字段召回、SQL 编译和执行结果。",
        "如果回答稳定且需要复用，可先保存到看板。",
        "如果需要给外部系统调用，再发布为接口；如果回答异常，进入查询审计排查 trace。",
    ])
    add_heading(doc, "3.3 结果确认", 2)
    add_bullets(doc, [
        "回答必须引用具体数据结果，而不是只有泛泛建议。",
        "SQL、图表和文字解释的指标口径应一致。",
        "如果系统给出候选数据集、候选字段或风险提示，用户需要先确认后再采纳结论。",
    ])
    doc.add_page_break()

    add_heading(doc, "4. 数据集 & 指标：维护语义能力和分析蓝图", 1)
    add_para(doc, "数据集页用于把物理表、字段、指标、维度、业务术语、分析蓝图和数据集路由契约组织成可被 Agent 使用的语义能力。这里不是单一页面，而是一组治理工作台，用户需要通过多个 Tab 分步完成。")
    add_screenshot(doc, "datasets", "截图 3：数据集语义能力与分析蓝图工作台")
    add_heading(doc, "4.1 进入与全局操作", 2)
    add_bullets(doc, [
        "左侧数据集列表用于切换当前维护对象，当前数据集会在页面标题和列表中高亮。",
        "顶部“导入 YAML / 导出 YAML”用于批量迁移或备份语义配置。",
        "顶部“同步表结构”用于从数据源刷新表和字段；字段变化后应先同步，再继续字段标注、指标、维度维护。",
        "顶部“AI 自动标注”用于让系统生成字段初标，生成后仍需要人工审核。",
        "顶部“新建数据集”用于创建新的语义治理单元。",
    ])
    add_heading(doc, "4.2 数据集主 Tab 功能", 2)
    add_detail_table(doc, ["Tab", "用途", "操作与验收"], [
        ("数据表", "决定哪些物理表进入当前数据集语义层。", "搜索表名并勾选需要纳入问数的表；点击表名查看字段和数据预览。完成后确认已选表数量正确，无关表没有被勾选。"),
        ("字段标注", "维护字段中文含义、字段角色和治理状态。", "按表、角色、来源筛选字段；点击 AI 自动标注；对候选指标、维度、时间字段、ID 字段逐条确认或忽略。完成后字段状态应变为已确认。"),
        ("指标", "沉淀可复用的度量口径。", "新建或编辑指标，填写名称、主表、表达式、时间字段和同义词；也可从候选字段快速创建。完成后问数链路应能召回该指标。"),
        ("维度", "沉淀可用于分组、筛选和关联的分析维度。", "新建或编辑维度，填写字段、表、关联字段、枚举值和同义词。完成后维度字段和关联关系应正确。"),
        ("分析蓝图", "沉淀复杂分析能力。", "点击新建蓝图，选择 SQL 导入或业务描述生成。发布前检查参数、输出列和测试；发布后确认状态、版本、触发词和最近验证时间。"),
        ("语义验证", "用自然语言问题验证路由、字段召回和 SQL 生成。", "保存验证问题并运行自检，观察数据集路由、资产命中、SQL 和失败原因。失败用例要回流到字段标注、词典或蓝图。"),
    ], [1150, 2850, 5360])
    add_heading(doc, "4.2.1 主 Tab 页面截图", 3)
    add_screenshot_series(doc, [
        ("dataset_tables_tab", "截图 3-1：数据表 Tab"),
        ("dataset_field_annotation_tab", "截图 3-2：字段标注 Tab"),
        ("dataset_metrics_tab", "截图 3-3：指标 Tab"),
        ("dataset_dimensions_tab", "截图 3-4：维度 Tab"),
        ("dataset_blueprints_tab", "截图 3-5：分析蓝图 Tab"),
        ("dataset_validation_tab", "截图 3-6：语义验证 Tab"),
    ])
    doc.add_page_break()
    add_heading(doc, "4.3 高级治理 Tab 功能", 2)
    add_detail_table(doc, ["Tab", "解决什么问题", "关键操作"], [
        ("SubAgent Manifest", "定义当前数据集作为 SubAgent 能力被 LeadAgent 路由时的契约。", "维护业务域、数据集描述、样例问题、负向样例和路由自检；保存草稿后再发布当前版本。"),
        ("语义词典", "维护业务术语、别名、口径说明和冲突检测。", "新增术语，绑定指标/维度/字段；使用冲突检测识别一词多义；对 AI 建议的别名进行人工采纳。"),
        ("分析场景", "组织业务场景级别的分析入口。", "把一组问题、蓝图和业务说明组合成场景，便于用户按业务主题使用。"),
        ("权限", "控制数据集和语义资产的可见与可用范围。", "设置哪些角色或成员可查看、维护、发布或调用当前数据集。"),
        ("版本历史", "追踪语义配置、Manifest 和蓝图的变更。", "查看历史版本、发布时间、变更说明；必要时用于回滚或审计。"),
    ], [1450, 3350, 4560])
    add_heading(doc, "4.3.1 高级治理页面截图", 3)
    add_screenshot_series(doc, [
        ("dataset_manifest_tab", "截图 3-7：SubAgent Manifest Tab"),
        ("dataset_terms_tab", "截图 3-8：语义词典 Tab"),
        ("dataset_scenarios_tab", "截图 3-9：分析场景 Tab"),
        ("dataset_permissions_tab", "截图 3-10：权限 Tab"),
        ("dataset_version_history_tab", "截图 3-11：版本历史 Tab"),
    ])
    add_heading(doc, "4.4 分析蓝图详情页 Tab", 2)
    add_detail_table(doc, ["详情 Tab", "用途", "用户需要检查的内容"], [
        ("概览", "查看蓝图基本信息、状态、来源、触发词和版本。", "确认业务描述、适用数据集、发布状态和最近验证时间是否正确。"),
        ("参数 · L1", "定义用户提问时可传入的参数。", "检查参数名、类型、必填项、默认值和自然语言映射方式。"),
        ("输出列 · L1", "定义结果表向用户展示哪些列。", "确认列名、中文名、格式、排序和是否适合对外展示。"),
        ("业务逻辑 · L2", "描述复杂过滤、分组、计算和业务规则。", "检查 SQL 片段、规则说明、风险提示和边界条件。"),
        ("测试", "用样例问题验证蓝图是否能稳定执行。", "输入测试问题，运行后检查 SQL、结果、错误提示和耗时。"),
        ("使用记录", "查看蓝图被触发和调用的历史。", "关注命中次数、失败记录、最近调用时间和用户反馈。"),
    ], [1500, 2700, 5160])
    add_screenshot(doc, "blueprint_detail_overview_tab", "截图 3-12：分析蓝图详情概览")
    add_heading(doc, "4.5 新建蓝图的两种方式", 2)
    add_detail_table(doc, ["方式", "流程", "适合场景"], [
        ("SQL 导入", "上传 SQL -> AI 审核 -> 精细配置 -> 触发与测试", "已有稳定 SQL，希望把它包装成可复用自然语言能力。"),
        ("业务描述生成", "描述场景 -> AI 草稿 -> 高级配置 -> 发布与测试", "只有业务分析诉求，需要系统先生成蓝图草稿，再由分析师审核。"),
    ], [1700, 3800, 3860])
    add_heading(doc, "4.6 推荐治理顺序", 2)
    add_steps(doc, [
        "先在“数据表”选择进入语义层的表。",
        "再到“字段标注”确认字段含义和角色。",
        "把高频度量沉淀到“指标”，把分组和筛选字段沉淀到“维度”。",
        "把业务别名和口径沉淀到“语义词典”。",
        "在“分析蓝图”中沉淀复杂流程。",
        "最后用“SubAgent Manifest”和“语义验证”检查路由、问法和执行稳定性。",
    ])
    add_callout(doc, "使用边界", "不要把所有表和字段一次性暴露给问数链路。语义资产越清晰，Agent 路由、SQL 生成和结果解释越稳定。", BLUE_GRAY_FILL)
    doc.add_page_break()

    add_heading(doc, "5. 数据源：接入数据库并同步表结构", 1)
    add_para(doc, "数据源页用于配置数据库连接、测试连接状态、同步 DDL/Schema，并为后续数据集治理提供物理表结构。数据源接入成功后，还需要到数据集页选择哪些表进入语义层。")
    add_screenshot(doc, "datasources", "截图 4：数据源管理页")
    add_heading(doc, "5.1 页面区域", 2)
    add_detail_table(doc, ["区域", "功能", "操作说明"], [
        ("顶部统计", "展示已连接数量、总表数、数据源数和平均延迟。", "先看这里判断连接状态是否整体正常。"),
        ("左侧数据源列表", "切换当前数据源并查看连接状态。", "点击某个数据源后，右侧展示该数据源详情。"),
        ("新建数据源", "新增数据库连接。", "填写类型、Host、端口、库名、账号、密码，保存前先测试连接。"),
        ("刷新状态", "重新拉取连接状态和统计。", "连接异常或同步后点击刷新，确认页面数据已更新。"),
    ], [1600, 3300, 4460])
    add_heading(doc, "5.2 数据源详情 Tab", 2)
    add_detail_table(doc, ["Tab", "具体功能", "怎么用", "成功标志"], [
        ("概览", "查看数据库类型、连接地址、表数量、驱动状态和基础配置。", "进入数据源后默认查看；如果页面提示驱动缺失，需要先补齐驱动或调整连接配置。", "状态显示已连接，库名、Host、端口和表数量符合预期。"),
        ("Schema", "查看 schema、表、字段、行数和大小。", "选择 schema，搜索表名；展开表查看字段类型、主键、是否为空等信息。", "能看到目标业务表和字段，字段类型与数据库一致。"),
        ("DDL 同步", "把数据库结构同步到语义层。", "点击同步按钮，等待进度完成；同步后回到数据集页继续数据表和字段治理。", "同步结果成功，更新时间刷新，数据集页能看到新表或新字段。"),
    ], [1200, 2500, 3560, 2100])
    add_heading(doc, "5.2.1 数据源详情 Tab 截图", 3)
    add_screenshot_series(doc, [
        ("datasource_overview_tab", "截图 4-1：数据源概览 Tab"),
        ("datasource_schema_tab", "截图 4-2：Schema Tab"),
        ("datasource_ddl_sync_tab", "截图 4-3：DDL 同步 Tab"),
    ])
    add_heading(doc, "5.3 接入流程", 2)
    add_steps(doc, [
        "点击“新建数据源”，选择数据库类型。",
        "填写连接参数，使用生产环境时优先使用只读账号。",
        "点击“测试连接”，确认网络、账号和权限都正常。",
        "保存后进入详情页，在 Schema Tab 检查目标表和字段。",
        "进入 DDL 同步 Tab 执行同步。",
        "回到“数据集 & 指标”，选择要进入语义层的表并继续字段标注。",
    ])
    doc.add_page_break()

    add_heading(doc, "6. 查询审计：回看 Trace、SQL、成本和失败原因", 1)
    add_para(doc, "查询审计页用于回看每一次问数的完整链路。它面向数据分析师、研发和运维人员，用于定位错误 SQL、错误路由、低置信度回答、成本异常和用户反馈。")
    add_screenshot(doc, "audit", "截图 5：查询审计页")
    add_heading(doc, "6.1 页面区域与功能", 2)
    add_detail_table(doc, ["区域", "功能", "怎么用"], [
        ("顶部指标", "查看查询总数、成功率、失败数、Token 和成本。", "先判断是否存在失败率升高、成本异常或 Token 暴涨。"),
        ("Trace 列表", "按时间查看问数记录。", "点击一条记录，右侧会显示该 trace 的问题、状态和执行链路。"),
        ("详情面板", "查看问题、回答预览、SQL、数据集、模型和入口路径。", "用于判断是问法问题、路由问题、SQL 问题还是执行问题。"),
        ("Observation 瀑布", "展示 Agent 节点执行顺序。", "逐个展开节点，检查输入、输出、耗时、错误和元数据。"),
        ("反馈与评分", "查看用户反馈、置信度和质量评分。", "把低分记录回流到语义词典、字段标注或蓝图测试用例。"),
    ], [1500, 2900, 4960])
    add_heading(doc, "6.2 常见排障动作", 2)
    add_detail_table(doc, ["现象", "优先检查", "处理方向"], [
        ("回答没有数据", "SQL 是否执行成功、过滤条件是否过窄、时间范围是否正确。", "放宽条件或修正时间字段，必要时补充维度和指标口径。"),
        ("选错数据集", "入口意图、数据集描述、Manifest 样例问题和负向样例。", "更新 SubAgent Manifest，增加正反例并重新验证。"),
        ("SQL 字段错误", "字段召回、字段标注、指标表达式、维度关联字段。", "回到数据集页修字段标注、指标或维度。"),
        ("成本异常", "模型调用次数、上下文长度、分析深度、重试次数。", "降低分析深度，优化 Prompt 或模型角色绑定。"),
        ("Langfuse 不可用", "页面是否展示本地 trace index 或 step_trace。", "使用本地回退链路继续排障，恢复 Langfuse 后再补查完整链路。"),
    ], [1700, 3760, 3900])
    add_heading(doc, "6.3 从聊天页跳转审计", 2)
    add_steps(doc, [
        "在对话问数页找到需要复核的回答。",
        "点击回答下方的 trace 或“查看链路”入口。",
        "系统带 trace_id 打开查询审计页。",
        "核对数据集、SQL、执行结果、风险提示和最终回答是否一致。",
    ])
    doc.add_page_break()

    add_heading(doc, "7. 查询历史：复用过往会话", 1)
    add_para(doc, "查询历史用于查找和继续过去的问数会话，适合复盘已完成分析、复用高频问题或继续多轮追问。页面会从后端实时拉取会话列表。")
    add_screenshot(doc, "history", "截图 6：查询历史页")
    add_heading(doc, "7.1 页面功能", 2)
    add_detail_table(doc, ["功能", "在哪里", "怎么用"], [
        ("搜索", "顶部搜索框", "输入问题、答案或 SQL 关键词，快速定位历史会话。"),
        ("时间筛选", "工具栏“近30天”", "缩小会话范围，适合排查近期问题或复盘最近分析。"),
        ("标签筛选", "工具栏“标签”", "按业务标签或问题类型筛选。"),
        ("状态 Tab", "全部、已发布接口、已收藏、草稿、已归档、我创建的、团队共享", "点击 Tab 切换列表范围，数字表示对应会话数量。"),
        ("继续会话", "会话列表行", "点击会话标题或整行进入 `/chat/{id}`，继续在原上下文追问。"),
        ("收藏", "行右侧书签按钮", "把高频或重要会话固定到已收藏筛选中。"),
        ("发布为接口", "行右侧 API 按钮", "把稳定会话转到 API 发布流程。"),
        ("删除", "行右侧删除按钮", "确认后删除会话；删除前应确认不再需要审计或复用。"),
        ("导出 / 新建文件夹", "页面右上角", "用于整理或外部分发历史分析材料。"),
    ], [1700, 3100, 4560])
    add_heading(doc, "7.2 使用建议", 2)
    add_bullets(doc, [
        "继续追问同一业务主题时使用历史会话，系统能复用上下文。",
        "完全切换业务主题时新建会话，避免旧上下文影响数据集路由。",
        "准备发布为接口前，先从历史会话确认问题、SQL 和结果是否稳定。",
    ])
    doc.add_page_break()

    add_heading(doc, "8. API 接口：把问数能力发布给系统调用", 1)
    add_para(doc, "API 接口页用于把稳定的问数能力发布给外部系统或业务流程调用。它把一次经过验证的问数会话固化为带参数、权限、日志和版本的服务接口。")
    add_screenshot(doc, "apis", "截图 7：API 接口页")
    add_heading(doc, "8.1 列表与状态筛选", 2)
    add_detail_table(doc, ["区域", "功能", "怎么用"], [
        ("顶部 KPI", "查看已发布接口数、24h 调用量、P95 延迟和成功率。", "先判断接口服务是否稳定，成功率和延迟是否在 SLO 内。"),
        ("从会话发布", "从稳定问数会话创建新接口。", "点击后选择来源会话，配置接口名称、路径、参数、返回格式和权限。"),
        ("API 文档", "查看调用规范。", "交付给外部系统前，先让调用方按文档确认路径、鉴权和参数。"),
        ("状态 Tab", "全部、已发布、测试中、草稿、已暂停。", "按接口生命周期筛选，重点关注测试中和已暂停接口。"),
        ("接口列表", "显示接口名称、路径、状态、版本、数据集和调用状态。", "点击某个接口后，右侧打开详情。"),
    ], [1600, 3200, 4560])
    add_heading(doc, "8.2 接口详情 Tab", 2)
    add_detail_table(doc, ["详情 Tab", "功能", "用户需要检查什么"], [
        ("概览", "查看接口描述、来源问题、数据集、调用趋势、配额、成功率和延迟。", "确认接口仍然可用，调用量、失败率和配额没有异常。"),
        ("参数", "维护发布时识别出的变量。", "检查参数类型、默认值、枚举、必填项和业务说明是否清楚。"),
        ("调用示例", "展示 HTTP 调用、SDK 或业务系统调用示例。", "交给调用方前先复制示例进行联调。"),
        ("调用日志", "查看调用时间、调用方、密钥、状态码、延迟和入参。", "排查外部调用失败、慢请求或权限问题。"),
        ("版本", "查看接口版本历史和变更说明。", "上线前确认是否有破坏性变更；必要时回滚到历史版本。"),
    ], [1450, 3000, 4910])
    add_heading(doc, "8.2.1 接口详情 Tab 截图", 3)
    add_screenshot_series(doc, [
        ("api_overview_tab", "截图 7-1：API 概览 Tab"),
        ("api_params_tab", "截图 7-2：参数 Tab"),
        ("api_examples_tab", "截图 7-3：调用示例 Tab"),
        ("api_logs_tab", "截图 7-4：调用日志 Tab"),
        ("api_versions_tab", "截图 7-5：版本 Tab"),
    ])
    add_heading(doc, "8.3 发布前检查", 2)
    add_bullets(doc, [
        "来源会话的 SQL 已在查询审计中复核。",
        "参数都有明确类型、默认值和业务含义。",
        "返回字段不包含不应对外暴露的敏感信息。",
        "接口权限、配额、调用方密钥和审计链路已经配置。",
    ])
    doc.add_page_break()

    add_heading(doc, "9. 系统设置：账号、安全、模型与集成配置", 1)
    add_para(doc, "系统设置页用于维护个人资料、通知、外观、工作区、成员权限、用量计费、数据源、LLM 模型、业务词典、API 密钥、Webhooks 和审计日志。左侧菜单按个人、工作区、数据与模型、开发者分组。")
    add_screenshot(doc, "settings", "截图 8：系统设置页")
    add_heading(doc, "9.1 个人设置", 2)
    add_detail_table(doc, ["菜单", "具体功能", "操作说明"], [
        ("账号与个人资料", "维护头像、姓名、邮箱、安全设置、密码、两步验证和活跃会话。", "个人信息变化后在这里更新；发现异常登录时先检查活跃会话并修改密码。"),
        ("通知", "配置邮件、站内、失败告警、周报等通知方式。", "管理员和分析师建议开启失败告警和高成本告警。"),
        ("外观", "配置主题、界面密度或显示偏好。", "根据使用习惯调整，不影响问数结果。"),
    ], [1700, 3600, 4060])
    add_screenshot_series(doc, [
        ("settings_profile", "截图 8-1：账号与个人资料"),
        ("settings_notifications", "截图 8-2：通知设置"),
        ("settings_appearance", "截图 8-3：外观设置"),
    ])
    add_heading(doc, "9.2 工作区设置", 2)
    add_detail_table(doc, ["菜单", "具体功能", "操作说明"], [
        ("工作区设置", "维护工作区名称、默认语言、默认时区和组织信息。", "影响团队成员看到的默认工作空间信息。"),
        ("成员与权限", "管理成员、角色、权限边界和数据集访问范围。", "新增成员后应分配最小必要权限，不要默认授予管理权限。"),
        ("用量 & 计费", "查看调用量、Token、成本、配额和计费状态。", "成本异常时结合查询审计定位高消耗 trace。"),
    ], [1700, 3600, 4060])
    add_screenshot_series(doc, [
        ("settings_workspace", "截图 8-4：工作区设置"),
        ("settings_members_permissions", "截图 8-5：成员与权限"),
        ("settings_usage_billing", "截图 8-6：用量与计费"),
    ])
    add_heading(doc, "9.3 数据与模型", 2)
    add_detail_table(doc, ["菜单", "具体功能", "操作说明"], [
        ("数据源", "查看或跳转数据源配置。", "数据库连接、Schema 和 DDL 同步仍建议在“数据源”主页面完成。"),
        ("LLM 模型", "配置 OpenAI-compatible / LiteLLM Proxy 模型。", "填写模型名称、Base URL、API Key 和状态；保存后测试连通性。"),
        ("模型角色绑定", "把 default、intent、dsl、sql_audit、report、annotation、blueprint 等任务绑定到具体模型。", "模型变更会影响问数质量，生产环境应先在测试环境验证。"),
        ("业务词典", "维护跨数据集通用术语和口径。", "与数据集内“语义词典”配合使用，避免同词不同义。"),
    ], [1700, 3600, 4060])
    add_screenshot_series(doc, [
        ("settings_data_sources", "截图 8-7：设置里的数据源"),
        ("settings_llm_models", "截图 8-8：LLM 模型"),
        ("settings_business_dictionary", "截图 8-9：业务词典"),
    ])
    add_heading(doc, "9.4 开发者设置", 2)
    add_detail_table(doc, ["菜单", "具体功能", "操作说明"], [
        ("API 密钥", "生成、停用和轮换外部调用密钥。", "密钥只展示一次，应按调用方分发独立密钥并定期轮换。"),
        ("Webhooks", "配置事件推送地址。", "用于把接口调用、失败告警或审计事件推送到外部系统。"),
        ("审计日志", "查看设置变更、密钥操作、权限调整和接口发布记录。", "排查权限或配置问题时先查审计日志。"),
    ], [1700, 3600, 4060])
    add_screenshot_series(doc, [
        ("settings_api_keys", "截图 8-10：API 密钥"),
        ("settings_webhooks", "截图 8-11：Webhooks"),
        ("settings_audit_logs", "截图 8-12：审计日志"),
    ])
    add_heading(doc, "9.5 保存与验证", 2)
    add_bullets(doc, [
        "表单保存后应出现明确的成功或失败反馈。",
        "模型配置保存后要执行连通性测试，并用一条真实问数验证效果。",
        "API 密钥、Webhook、成员权限等敏感变更应能在审计日志中追踪。",
        "编辑模型时 API Key 留空通常表示保留旧密钥，不要把密钥明文复制到文档或聊天中。",
    ])
    doc.add_page_break()

    add_heading(doc, "10. 常见操作检查清单", 1)
    add_heading(doc, "10.1 普通问数用户", 2)
    add_bullets(doc, [
        "问题中包含对象、时间、指标和拆分维度。",
        "回答后检查 SQL、图表、结果表和风险提示。",
        "回答不可信时点击“查看链路”或请数据分析师从查询审计排查。",
        "需要继续分析时优先在同一会话追问；切换业务主题时新建对话。",
    ])
    add_heading(doc, "10.2 数据分析师", 2)
    add_bullets(doc, [
        "新数据源接入后先同步表结构，再创建或更新数据集。",
        "字段标注、指标、维度和语义词典需要人工复核。",
        "高频复杂分析应沉淀为分析蓝图，并维护触发词和测试用例。",
        "定期查看查询审计中的失败 trace，把问题回流到语义治理。 ",
    ])
    add_heading(doc, "10.3 管理员/运维", 2)
    add_bullets(doc, [
        "确保数据源账号权限最小化，优先只读。",
        "确保 default、dsl、report 等核心模型角色已绑定可用模型。",
        "关注查询审计里的失败率、token、成本和高延迟记录。",
        "外部发布接口前确认权限、参数、返回字段和审计链路。",
    ])

    doc.core_properties.title = "数语多智能体产品用户手册"
    doc.core_properties.subject = "带真实页面截图的数语操作手册"
    doc.core_properties.author = "yangkai"
    doc.core_properties.comments = "由 scripts/export_user_operation_manual_docx.py 生成"
    doc.core_properties.created = datetime(2026, 6, 12, 18, 20, 0)
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_manual()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
