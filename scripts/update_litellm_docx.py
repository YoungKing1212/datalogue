# ============================================================
# File Name   : update_litellm_docx.py
# Description:
#   复现 2026-06-10 LiteLLM 历史设计方案 DOCX。
#
# Responsibilities:
#   - 基于归档的 20260605 版本生成当时的 20260610 更新版。
#   - 只写入历史归档目录，不作为当前产品规划入口。
#
# Author      : yangkai
# Created On  : 2026-06-10
# ============================================================

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARCHIVE_DIR = ROOT / "assets/archive/2026-07-03-legacy-design-plans"
# 该脚本仅用于复现历史文档，禁止把输出重新标记为当前架构方案。
SRC = ARCHIVE_DIR / "Datalogue_设计开发方案_v2.0_更新版_20260605.docx"
OUT = ARCHIVE_DIR / "Datalogue_设计开发方案_v2.0_更新版_20260610.docx"


def text_of(paragraph) -> str:
    return " ".join(paragraph.text.split())


def paragraph_after(paragraph, text: str = "", style: str | None = None):
    """在指定段落后插入新段落，尽量复用原文档样式。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraphs_after(doc: Document, anchor_text: str, items: list[tuple[str, str | None]]) -> None:
    for para in doc.paragraphs:
        if text_of(para) == anchor_text:
            current = para
            for text, style in items:
                current = paragraph_after(current, text, style)
            return
    raise RuntimeError(f"找不到锚点段落: {anchor_text}")


def set_paragraph_text(doc: Document, old_text: str, new_text: str) -> None:
    for para in doc.paragraphs:
        if text_of(para) == old_text:
            para.clear()
            para.add_run(new_text)
            return
    raise RuntimeError(f"找不到待替换段落: {old_text}")


def add_table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        cell.text = value


def set_cell_text(cell, value: str) -> None:
    cell.text = value


def shade_header_rows(doc: Document) -> None:
    """保持新增行不会破坏原表格；仅补齐空白表头底色。"""
    for table in doc.tables:
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), "EAF1F8")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(15, 23, 42)


def main() -> None:
    doc = Document(SRC)

    # 封面更新说明
    set_paragraph_text(
        doc,
        "（v2.0.1 更新：新增分析蓝图一等语义资产 · 对齐当前实现 · 补充差距矩阵与飞书任务优先级）",
        "（v2.0.2 更新：新增 LiteLLM 前端配置化接入 · 数据库持久化模型配置 · 按任务角色绑定 LLM）",
    )

    # 技术选型表：LLM 接入行
    tech_table = doc.tables[1]
    tech_table.rows[7].cells[1].text = "LiteLLM Proxy / OpenAI-compatible 网关 / ChatOpenAI 客户端"
    tech_table.rows[7].cells[2].text = (
        "前端配置模型来源并加密落库；按 default/intent/dsl/sql_audit/report/"
        "annotation/blueprint 角色路由；保留 LangGraph 流式事件兼容。"
    )

    # 双模部署表：LLM 接入行
    deploy_table = doc.tables[3]
    deploy_table.rows[2].cells[1].text = (
        "优先内网 LiteLLM Proxy 或企业模型网关，统一 OpenAI-compatible 协议；"
        "供应商密钥可企业自管。"
    )
    deploy_table.rows[2].cells[2].text = (
        "可接闭源大模型 API 或托管 LiteLLM Proxy，按租户/角色绑定模型与配额。"
    )

    # 4.4 正文替换与补充
    set_paragraph_text(
        doc,
        "【功能描述】支持配置多个 LLM 来源（Claude/GPT-4o/Qwen/DeepSeek 等），按任务类型路由到合适模型，并在故障/限流时自动降级。",
        (
            "【功能描述】支持通过前端维护多个 OpenAI-compatible / LiteLLM Proxy 模型来源，"
            "并把模型配置持久化到数据库。系统按任务角色选择模型：默认、意图识别、"
            "DSL/SQL 生成、SQL 审计、报告生成、字段标注和分析蓝图。"
        ),
    )
    set_paragraph_text(
        doc,
        "【解决方案】LLM 来源 CRUD（endpoint、模型、密钥、配额、价格）；路由策略按节点任务分配——意图识别等轻任务用低成本模型，DSL 生成等关键任务用强模型；故障/超时/限流时按优先级链路自动 fallback；私有化下可全部指向内网模型。",
        (
            "【解决方案】新增 llm_model_config 与 llm_role_binding 两张表，前端系统设置页提供模型"
            "新增、编辑、启停、删除、测试连接和角色绑定。API Key 通过 AES 加密写入 api_key_enc，"
            "接口响应只返回 api_key_set，不回传明文。运行时 get_llm(role, db) 先读取角色绑定模型，"
            "未配置时回退 default，再回退 .env 中 OPENAI_API_KEY / OPENAI_BASE_URL / LLM_MODEL。"
        ),
    )
    set_paragraph_text(
        doc,
        "【解决的问题】解决「单一模型成本高/有故障风险/私有化不可用闭源 API」。按任务路由兼顾准确率与成本，降级保障可用性，双模部署下灵活切换。",
        (
            "【解决的问题】解决「模型供应商绑定在环境变量、无法在页面调整、不同任务无法选用不同模型」的问题。"
            "保留 ChatOpenAI 作为 LangChain OpenAI-compatible 客户端，避免破坏现有 invoke/astream、"
            "LangGraph astream_events 和 SSE token 流，同时允许底层实际接入 LiteLLM Proxy、Qwen、DeepSeek、"
            "OpenAI 或企业私有网关。"
        ),
    )
    insert_paragraphs_after(
        doc,
        (
            "【解决的问题】解决「模型供应商绑定在环境变量、无法在页面调整、不同任务无法选用不同模型」的问题。"
            "保留 ChatOpenAI 作为 LangChain OpenAI-compatible 客户端，避免破坏现有 invoke/astream、"
            "LangGraph astream_events 和 SSE token 流，同时允许底层实际接入 LiteLLM Proxy、Qwen、DeepSeek、"
            "OpenAI 或企业私有网关。"
        ),
        [
            (
                "【当前实现状态（2026-06-10）】已完成前端配置化 MVP：/api/llm/models 提供模型配置 CRUD，"
                "/api/llm/models/{id}/test 提供连接测试，/api/llm/role-bindings 保存任务角色绑定；"
                "问数链路、字段标注和分析蓝图已按角色调用 get_llm。未完成部分为多级 fallback 链、"
                "按模型维度的成本/Token 审计、租户级隔离和密钥轮换审计。",
                None,
            ),
            (
                "【验收结果】后端全量测试 219 passed、4 skipped；重点回归 128 passed、3 skipped；"
                "Alembic 当前 head 为 e3f4a5b6c7d8；前端 lint 0 errors、build 通过；设置页 LLM 模型区已完成渲染检查。",
                None,
            ),
        ],
    )

    # 数据库表设计：LLM/发布/评估
    db_table = doc.tables[24]
    db_table.rows[1].cells[0].text = "llm_model_config"
    db_table.rows[1].cells[1].text = "表"
    db_table.rows[1].cells[2].text = (
        "LLM 模型连接配置：name、provider、base_url、model、api_key_enc、status、"
        "request_timeout_seconds、last_test_result"
    )
    db_table.rows[2].cells[0].text = "llm_role_binding"
    db_table.rows[2].cells[1].text = "表"
    db_table.rows[2].cells[2].text = (
        "任务角色绑定：role → model_config_id；内置 default、intent、dsl、sql_audit、"
        "report、annotation、blueprint"
    )
    add_table_row(
        db_table,
        [
            "api_endpoint",
            "表",
            "已发布 API：path、dataset_id、sql_template、param_schema(JSON)、version、status",
        ],
    )

    # API 表：语义与配置类追加 LLM 管理接口
    api_table = doc.tables[26]
    add_table_row(
        api_table,
        [
            "GET/POST",
            "/api/llm/models",
            "LLM 模型配置",
            "前端维护 OpenAI-compatible / LiteLLM 模型来源，API Key 加密保存且不回显",
        ],
    )
    add_table_row(
        api_table,
        [
            "PUT",
            "/api/llm/role-bindings",
            "LLM 角色绑定",
            "按 default/intent/dsl/sql_audit/report/annotation/blueprint 绑定启用模型",
        ],
    )
    add_table_row(
        api_table,
        [
            "POST",
            "/api/llm/models/{id}/test",
            "模型连通性测试",
            "调用模型网关做轻量测试并保存最近一次测试结果",
        ],
    )

    # 路线图 W11：标注已完成 MVP 与剩余项
    roadmap_table = doc.tables[32]
    roadmap_table.rows[4].cells[1].text = (
        "LLM 多来源配置：前端配置、加密落库、角色绑定已完成 MVP；后续补故障降级、成本审计和租户级策略。"
    )
    roadmap_table.rows[4].cells[2].text = (
        "模型配置管理可用；角色路由已接入；降级链路待增强。"
    )

    # 差距矩阵：把语义验证和运行时更新到现状，并追加 LLM 行
    gap_table = doc.tables[37]
    gap_table.rows[6].cells[2].text = (
        "已增强：语义验证覆盖入口路由、普通问数路径、术语命中、蓝图命中、生成 SQL、失败原因、置信度风险和节点轨迹；支持保存最近验证用例。"
    )
    gap_table.rows[6].cells[3].text = (
        "批量回放调度、断言规则和定时评测仍未实现。"
    )
    gap_table.rows[7].cells[2].text = (
        "已有：确定性入口分类、术语冲突澄清、语义资产解析、DSL/SQL 生成、SQL 审计修复、报告生成与 SSE 事件。"
    )
    gap_table.rows[7].cells[3].text = (
        "LeadAgent 复杂拆解、WeKnora 运行时检索、反馈学习闭环仍未完整落地。"
    )
    add_table_row(
        gap_table,
        [
            "LLM 多来源配置",
            "前端维护多模型来源，按任务角色路由到合适模型，并支持 LiteLLM Proxy / 私有网关。",
            "已完成：llm_model_config、llm_role_binding、/api/llm、设置页模型管理、角色绑定、连接测试、数据库配置优先和 .env 兜底。",
            "故障降级链、模型用量/成本审计、租户隔离、密钥轮换和真实网关联调仍待补齐。",
            "P1",
        ],
    )

    # 后续任务：追加 LLM 后续闭环任务
    task_table = doc.tables[38]
    add_table_row(
        task_table,
        [
            "T-014",
            "P1",
            "LLM 故障降级与用量审计",
            "在现有前端配置化基础上，为每个角色支持主备模型、超时/限流 fallback，并记录模型、Token、耗时、成本和错误码。",
            "角色可配置 fallback 顺序；模型失败自动切换并写审计；查询历史可看到 role/model/token/latency；测试覆盖成功、失败、降级和无配置兜底。",
        ],
    )

    shade_header_rows(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
