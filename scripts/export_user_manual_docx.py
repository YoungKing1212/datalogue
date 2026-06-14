# ============================================================
# File Name   : export_user_manual_docx.py
# Description:
#   生成“数语”多智能体产品用户手册 DOCX，供交付、培训和内部验收使用。
#
# Responsibilities:
#   - 按固定 Word 样式生成用户手册正文、步骤、表格、页眉页脚和执行链路图。
#   - 将当前产品的多智能体问数、数据源、语义治理、分析蓝图、审计和系统设置能力整理成可阅读手册。
#
# Author      : yangkai
# Created On  : 2026-06-12
# ============================================================

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "docs" / "数语多智能体产品用户手册.docx"
FLOW_IMAGE = ROOT / "docs" / "datalogue_full_execution_flow.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(86, 100, 118)
GRAY_FILL = "F2F4F7"
BLUE_GRAY_FILL = "E8EEF5"
LIGHT_FILL = "F8FAFC"
BORDER = "D7DEE8"


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: RGBColor | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths: list[int], header: bool = True) -> None:
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, color=INK)
        if row_idx == 0 and header:
            for cell in row.cells:
                set_cell_shading(cell, BLUE_GRAY_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=9.5, color=INK, bold=True)


def next_numbering_id(numbering, tag_name: str, attr_name: str) -> int:
    values: list[int] = []
    for node in numbering.findall(qn(tag_name)):
        raw = node.get(qn(attr_name))
        if raw and raw.isdigit():
            values.append(int(raw))
    return (max(values) + 1) if values else 1


def new_decimal_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(tabs)
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_para(doc: Document, text: str = "", *, style: str | None = None,
             size: float = 11, bold: bool = False, color: RGBColor = INK,
             after: float | None = None, before: float | None = None,
             align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_steps(doc: Document, items: list[str]) -> None:
    num_id = new_decimal_numbering_id(doc)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=INK)
    add_para(doc, "", after=4)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    for level, before, after, size, color in [
        (1, 18, 10, 16, BLUE),
        (2, 14, 7, 13, BLUE),
        (3, 10, 5, 12, DARK_BLUE),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("数语多智能体产品用户手册")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Datalogue / 内部用户手册")
    set_run_font(r, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    add_para(doc, "产品手册", size=11, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=18)
    add_para(doc, "数语多智能体产品用户手册", size=28, bold=True, color=DARK_BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "AI 原生智能问数、语义治理、分析蓝图与查询审计使用指南", size=14,
             color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)

    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("适用对象", "终端问数用户、数据分析师、数据管理员、交付和运维人员"),
        ("适用版本", "以 2026-06-12 本地多智能体工作区能力为基准"),
        ("产品范围", "LeadAgent 控制面、Dataset SubAgent 数据面、数据源、数据集、分析蓝图、审计和模型设置"),
        ("输出格式", "Microsoft Word DOCX"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
    style_table(table, [1800, 7560], header=False)
    for row in table.rows:
        set_cell_shading(row.cells[0], BLUE_GRAY_FILL)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=10, bold=True, color=DARK_BLUE)
    add_para(doc, "", after=54)
    add_para(doc, "生成日期：2026-06-12", size=10.5, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "维护人：yangkai / KenYang", size=10.5, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    sections = [
        "1. 手册定位与产品边界",
        "2. 角色与常见工作流",
        "3. 登录后界面导航",
        "4. 数据源管理",
        "5. 数据集与语义治理",
        "6. 分析蓝图工作台",
        "7. 对话问数与多轮分析",
        "8. 查询审计与链路回看",
        "9. 系统设置与模型配置",
        "10. 常见问题与排障",
        "11. 术语表",
        "12. 附录：上线前检查清单",
    ]
    for item in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)
    doc.add_page_break()


def add_role_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    headers = ["角色", "核心目标", "主要入口", "典型产出"]
    for idx, text in enumerate(headers):
        table.cell(0, idx).text = text
    rows = [
        ("终端问数用户", "用自然语言查询数据并获得解释", "对话问数、历史、收藏", "回答、SQL、图表、追问建议"),
        ("数据分析师", "维护指标、维度、术语和蓝图", "数据集 & 指标、分析蓝图、语义验证", "可复用语义资产、验证用例"),
        ("数据管理员", "接入数据源并控制可见范围", "数据源、数据集、权限", "连接配置、表结构、权限策略"),
        ("运维/研发", "定位失败链路、成本和质量问题", "查询审计、系统设置、Langfuse Trace", "Trace、token、失败原因、反馈记录"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1700, 2500, 2500, 2660])


def build_manual() -> Document:
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc.sections[0])
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "1. 手册定位与产品边界", 1)
    add_para(doc, "数语是一套面向企业数据分析场景的 AI 原生智能问数平台。用户用自然语言提出业务问题，系统通过 LeadAgent 进行问题理解、澄清、路由和多轮上下文管理，再将明确的数据集任务交给 Dataset SubAgent 执行语义召回、DSL/SQL 生成、SQL 执行和报告生成。")
    add_callout(doc, "产品边界", "当前手册聚焦“场景 2：Agent 驱动智能问数”。它不是传统报表系统，也不是只维护固定 SQL 模板的低代码查询工具；分析蓝图、语义词典和数据集 Manifest 都服务于让 Agent 更可靠地理解企业语义。", BLUE_GRAY_FILL)
    add_bullets(doc, [
        "面向自然语言问数：支持普通指标查询、明细查询、分析蓝图命中、澄清追问和拒答保护。",
        "面向语义治理：通过数据表、字段标注、指标、维度、业务术语、Manifest 和验证用例提高问数质量。",
        "面向生产审计：通过查询审计、Trace、Scores、token 和成本信息回看每次问数链路。",
        "面向多轮交互：通过 ConversationStore 和 Capsule 记录当前会话的数据集、查询上下文和结果摘要，支撑继续追问、切换主题和结果解释。",
    ])

    add_heading(doc, "2. 角色与常见工作流", 1)
    add_role_table(doc)
    add_heading(doc, "2.1 终端问数用户快速路径", 2)
    add_steps(doc, [
        "进入“对话问数”，选择或沿用当前数据集。",
        "用业务语言提出问题，例如“查看去年杨凯的日报情况”或“最近 30 天各部门任务完成趋势”。",
        "观察右侧或消息中的推理过程、数据集命中、SQL、结果表和最终解释。",
        "如果系统要求澄清，按候选数据集、术语或条件点击确认，再继续追问。",
        "对回答进行点赞或点踩；发现口径错误时在反馈中说明原因，便于后续标注和回归。",
    ])
    add_heading(doc, "2.2 数据分析师治理路径", 2)
    add_steps(doc, [
        "在“数据源”完成连接和 Schema 同步。",
        "在“数据集 & 指标”选择物理表，完成字段标注、指标、维度和语义词典维护。",
        "对高频复杂分析创建分析蓝图，审核触发词、参数、输出列和业务步骤。",
        "在“语义验证”中用真实问法验证路由、资产召回、DSL 和 SQL 生成结果。",
        "在“查询审计”中查看低分、失败或高成本 trace，反向补齐术语、指标、维度和蓝图。",
    ])

    add_heading(doc, "3. 登录后界面导航", 1)
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["导航入口", "用途", "使用频率"]):
        table.cell(0, idx).text = text
    rows = [
        ("工作台", "查看总体入口和待处理事项。", "日常"),
        ("对话问数", "自然语言问数、多轮追问、查看推理过程、发布接口。", "高"),
        ("数据集 & 指标", "维护表资产、字段标注、指标、维度、蓝图、Manifest、术语、权限和版本。", "高"),
        ("数据源", "新增数据库连接、测试连接、同步 Schema、预览表结构。", "中"),
        ("查询审计", "查看 trace 列表、成功率、失败数、token、成本和 observation 瀑布。", "中"),
        ("系统设置", "维护 LLM 模型配置、角色绑定和环境兜底。", "低到中"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1900, 5600, 1860])

    add_heading(doc, "4. 数据源管理", 1)
    add_para(doc, "数据源是问数能力的底层入口。只有连接、同步并选择进入数据集的表，才会进入后续语义治理和问数链路。")
    add_heading(doc, "4.1 新建数据源", 2)
    add_steps(doc, [
        "进入“数据源”。",
        "点击“新建数据源”，填写名称、数据库类型、Host、端口、库名、用户名和密码。",
        "保存前点击“测试连接”，确认网络、账号和权限可用。",
        "保存后进入详情页，查看连接状态、表数量和 Schema 信息。",
    ])
    add_heading(doc, "4.2 DDL 与 Schema 同步", 2)
    add_bullets(doc, [
        "“Schema”页用于查看当前已扫描到的库、表和字段。",
        "“DDL 同步”用于把数据库结构同步到数语语义层，适合表结构变更后执行。",
        "同步后需要回到数据集页面，选择哪些表纳入当前数据集，避免把无关或敏感表暴露给问数链路。",
    ])
    add_callout(doc, "安全提示", "数据源账号建议使用只读账号；涉及生产库时应控制可访问 schema、表和行列权限。数语的 SQL 生成链路会做只读和限制约束，但底层数据库权限仍是第一道边界。")

    add_heading(doc, "5. 数据集与语义治理", 1)
    add_para(doc, "数据集把一组业务相关表、指标、维度、术语、蓝图和路由契约组织在一起。多智能体问数中，LeadAgent 先选择合适的数据集，Dataset SubAgent 再在数据集内部执行问数。")
    add_heading(doc, "5.1 数据集能力页签", 2)
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["页签", "解决的问题", "关键操作"]):
        table.cell(0, idx).text = text
    rows = [
        ("数据表", "哪些物理表进入当前语义层。", "搜索表、勾选表、预览数据、同步表结构。"),
        ("字段标注", "字段表示什么业务含义。", "审核 AI 标注、设置字段角色、忽略无关字段。"),
        ("指标", "企业常用计算口径。", "新增指标、维护聚合规则、绑定字段。"),
        ("维度", "分析拆分口径和值域解释。", "新增维度、维护枚举说明、绑定字段。"),
        ("分析蓝图", "高频复杂分析路径。", "从 SQL 或业务描述生成蓝图、审核、测试、发布。"),
        ("语义验证", "问法是否能稳定命中正确资产。", "输入真实问题，检查路由、DSL、SQL 和命中资产。"),
        ("SubAgent Manifest", "数据集被 LeadAgent 选择的契约。", "维护路由问法、版本、自检样例和审核状态。"),
        ("语义词典", "跨资产别名、术语冲突和口径解释。", "维护业务术语、别名、定义和冲突处理。"),
        ("权限/版本历史", "治理可见范围和变更记录。", "后续按企业权限和发布流程完善。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1900, 3500, 3960])

    add_heading(doc, "5.2 查询约束", 2)
    add_para(doc, "数据集右键菜单中可维护查询约束。默认约束包括未指定时间时查最近 30 天、默认 LIMIT 100、最大 LIMIT 1000。这些约束会进入 Schema prompt、真实 Schema SQL 生成、推断 SQL 生成、语义 DSL 生成和 DSL 编译路径。")
    add_callout(doc, "约束边界", "查询约束主要约束 AI 生成和 DSL 编译路径；用户手写 SQL、历史已保存 SQL 或分析蓝图 SQL 模板仍需要在执行前结合 SQL 审计和数据库权限共同控制。")

    add_heading(doc, "6. 分析蓝图工作台", 1)
    add_para(doc, "分析蓝图用于把复杂 SQL、存储过程逻辑或手工业务步骤沉淀为可触发的问数能力。它适合经营日报、库存预警、客户价值、销售漏斗、NPS 归因等跨表、跨步骤的高频分析。")
    add_heading(doc, "6.1 创建方式", 2)
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["方式", "适用场景", "产出"]):
        table.cell(0, idx).text = text
    rows = [
        ("SQL 导入", "已有可执行 SQL 或可改造成 SELECT/WITH 的分析逻辑。", "SQL 模板蓝图、参数、输出列、触发词和步骤。"),
        ("手工业务描述", "业务路径清楚，但还没有稳定 SQL。", "semantic_plan 蓝图，作为问数链路的业务约束和计划上下文。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1900, 3800, 3660])
    add_heading(doc, "6.2 从草稿到发布", 2)
    add_steps(doc, [
        "点击“新建蓝图”，选择 SQL 导入或手工业务描述。",
        "让 AI 拆解蓝图草稿，生成名称、描述、触发问法、参数、输出列和业务步骤。",
        "人工审核自然语言字段，避免把 SQL 参数名、变量值或不适合用户看的技术细节暴露出去。",
        "在详情页查看概览、参数、输出列、业务逻辑、测试和使用记录。",
        "可选执行测试，确认 SQL 预览、结果数据和错误信息。",
        "发布蓝图。发布不强制要求测试成功，但未测试或失败蓝图应在交付前重点复核。",
    ])

    add_heading(doc, "7. 对话问数与多轮分析", 1)
    add_para(doc, "对话问数是终端用户最主要入口。系统会把用户问题转成多智能体控制面和数据面任务：LeadAgent 负责理解、澄清、路由、工具选择和最终叙述；Dataset SubAgent 负责数据集内部的语义召回、DSL/SQL、执行和结果摘要。")
    if FLOW_IMAGE.exists():
        add_para(doc, "图：数语问数执行链路", size=10, bold=True, color=MUTED, after=4)
        doc.add_picture(str(FLOW_IMAGE), height=Inches(4.85))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "7.1 发起问题", 2)
    add_steps(doc, [
        "进入“对话问数”或历史会话。",
        "输入业务问题，尽量包含对象、时间、指标和拆分维度。",
        "查看系统是否自动命中数据集、分析蓝图、业务术语或指标维度。",
        "如果右侧“推理过程”关闭，可点击顶部按钮打开，查看执行步骤。",
    ])
    add_heading(doc, "7.2 多轮追问", 2)
    add_bullets(doc, [
        "继续追问：例如“再看上个月”“按部门拆一下”，系统会尝试继承上一轮数据集和查询上下文。",
        "切换主题：如果问题明显换到另一个业务域，LeadAgent 会重新走数据集路由，不继承旧数据集。",
        "解释结果：例如“为什么这么高”，系统可基于上一轮 ResultDigest 做结果解释，而不是重新生成查询。",
        "闲聊或无关问题：LeadAgent 可在控制面停止，不调度 Dataset SubAgent。",
    ])
    add_heading(doc, "7.3 查看回答中的关键证据", 2)
    add_bullets(doc, [
        "SQL：用于核对系统实际查询了什么。",
        "结果表或图表：用于核对行数、维度和数值范围。",
        "推理过程：用于查看 DSL 校验、SQL 编译、SQL 执行和报告生成状态。",
        "Trace ID：用于进入查询审计页回看完整链路。",
        "置信度和风险提示：用于识别口径不确定、数据质量不足或结果需要人工复核的情况。",
    ])

    add_heading(doc, "8. 查询审计与链路回看", 1)
    add_para(doc, "查询审计用于把每次问数从用户问题、数据集、路由、Prompt、LLM 调用、DSL、SQL、执行结果、最终回答到用户反馈串起来。当前产品优先在 Datalogue 内部渲染 trace，不要求客户外跳 Langfuse 控制台。")
    add_heading(doc, "8.1 审计页指标", 2)
    add_bullets(doc, [
        "查询总数：当前范围内记录的 trace 数量。",
        "成功率和失败数：用于发现近期质量波动。",
        "Token 和成本：用于观察模型调用规模和成本趋势。",
        "Trace 列表：显示问题、答案预览、SQL 预览、入口路径和创建时间。",
        "Observation 瀑布：展开每个节点的输入、输出、元数据、耗时和模型信息。",
        "Scores：展示用户反馈或自动评测分数。",
    ])
    add_heading(doc, "8.2 失败排查顺序", 2)
    add_steps(doc, [
        "先看失败节点：是路由失败、DSL 校验失败、SQL 编译失败、SQL 执行失败，还是报告生成失败。",
        "再看输入上下文：数据集是否正确、时间范围是否继承、术语是否命中。",
        "核对 SQL：是否访问了错误表、字段、指标、维度或时间条件。",
        "检查数据源状态和权限：连接是否可用、账号是否有读权限、Schema 是否已经同步。",
        "把根因反馈到治理资产：补术语、补指标维度、修蓝图、加语义验证用例或调整查询约束。",
    ])

    add_heading(doc, "9. 系统设置与模型配置", 1)
    add_para(doc, "系统支持 OpenAI-compatible 协议的模型配置，也可以通过 LiteLLM Proxy 或私有模型网关接入不同供应商模型。模型配置优先级为：前端系统设置中的角色绑定模型、default 角色绑定模型、`.env` 兜底配置。")
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["模型角色", "用途", "配置建议"]):
        table.cell(0, idx).text = text
    rows = [
        ("default", "未单独绑定角色时的默认模型。", "先配置，保证系统有兜底。"),
        ("intent", "意图理解与入口路由。", "优先选择速度快、分类稳定的模型。"),
        ("dsl", "DSL/SQL 生成。", "优先选择结构化输出和 SQL 能力强的模型。"),
        ("sql_audit", "SQL 失败诊断。", "需要理解数据库方言和错误信息。"),
        ("report", "最终报告解释和流式回答。", "优先选择中文表达稳定的模型。"),
        ("annotation", "表字段自动标注。", "需要业务语义理解和字段解释能力。"),
        ("blueprint", "蓝图 SQL 草稿和业务场景理解。", "需要长上下文和结构化输出能力。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, [1800, 3400, 4160])
    add_callout(doc, "密钥保护", "前端写入的 API Key 会由后端加密保存，接口响应只返回是否已设置，不回传明文。编辑模型时 API Key 留空不会覆盖旧密钥。", BLUE_GRAY_FILL)

    add_heading(doc, "10. 常见问题与排障", 1)
    faq = [
        ("问数没有结果", "先确认数据源连接正常、数据集已选择表、字段已标注，再进入查询审计查看失败节点和 SQL 错误。"),
        ("系统要求选择数据集", "说明 LeadAgent 无法稳定判断业务域。请选择候选数据集，随后建议补充 Manifest 路由问法和语义验证用例。"),
        ("回答数值不对", "核对 SQL、指标口径、时间范围和维度拆分；若字段已被召回但 SQL 仍错，优先检查字段映射和 DSL 编译消费路径。"),
        ("蓝图发布后执行失败", "发布不等于测试通过。进入蓝图详情的测试页核对参数、SQL 预览、只读限制和数据库错误。"),
        ("Trace 明细不完整", "Langfuse 未启用或不可用时，系统会降级显示本地索引和 step_trace；完整 observation、token 和 cost 依赖有效 Langfuse 配置。"),
        ("多轮追问继承错了数据集", "明确说出新的业务域或数据集，必要时新建会话；数据团队应补充 Manifest 自检样例并复核 active dataset 状态。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "问题"
    table.cell(0, 1).text = "处理建议"
    for q, a in faq:
        cells = table.add_row().cells
        cells[0].text = q
        cells[1].text = a
    style_table(table, [2500, 6860])

    add_heading(doc, "11. 术语表", 1)
    terms = [
        ("LeadAgent", "多智能体控制面，负责理解问题、澄清、路由、工具选择、多轮状态和最终叙述。"),
        ("Dataset SubAgent", "数据集内部执行智能体，负责语义召回、DSL/SQL、执行和结果摘要。"),
        ("ConversationStore", "会话状态存储，用于保存 active dataset、capsule、轮次锁和多轮上下文。"),
        ("Capsule", "跨轮压缩状态，记录查询上下文、结果摘要、时间语义和下一轮可复用信息。"),
        ("Manifest", "数据集被 LeadAgent 识别和选择的契约，包含路由问法、自检样例、版本和审核状态。"),
        ("分析蓝图", "将复杂分析路径固化为可复用问数能力的资产，可以来自 SQL 或业务场景描述。"),
        ("DSL", "面向语义层的结构化查询表达，连接自然语言、语义资产和最终 SQL。"),
        ("Trace", "一次问数的全链路记录，包含节点、Prompt、LLM、SQL、结果、成本和反馈。"),
        ("Score", "用户反馈或自动评测分数，用于质量闭环。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "术语"
    table.cell(0, 1).text = "说明"
    for term, desc in terms:
        cells = table.add_row().cells
        cells[0].text = term
        cells[1].text = desc
    style_table(table, [2200, 7160])

    add_heading(doc, "12. 附录：上线前检查清单", 1)
    checklist = [
        "数据源连接测试通过，账号为只读或受控权限。",
        "目标数据集已选择必要表，移除了无关或敏感表。",
        "关键字段已完成人工复核，指标和维度口径清晰。",
        "高频复杂问题已沉淀为分析蓝图，并完成至少一次试运行。",
        "Manifest 路由问法和语义验证用例覆盖主要业务场景。",
        "系统设置中 default、dsl、report 等核心模型角色已绑定。",
        "查询审计页能看到真实 trace 列表和详情，失败时有本地 fallback。",
        "用户反馈入口可用，点踩问题有后续治理流程。",
        "多轮功能开启前已验证 continue、switch、interpret 和 chitchat 的边界。",
    ]
    add_bullets(doc, checklist)

    doc.core_properties.title = "数语多智能体产品用户手册"
    doc.core_properties.subject = "Datalogue 用户手册"
    doc.core_properties.author = "yangkai"
    doc.core_properties.comments = "由 scripts/export_user_manual_docx.py 生成"
    doc.core_properties.created = datetime(2026, 6, 12, 18, 0, 0)
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_manual()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
