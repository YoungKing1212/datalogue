# ============================================================
# File Name   : update_design_plan_v2.py
# Description:
#   复现 Datalogue 2026-06-05 历史设计方案的辅助脚本。
#
# Responsibilities:
#   - 基于归档的 v2.0 文档生成当时的修订版本。
#   - 只写入历史归档目录，不作为当前产品规划入口。
#
# Author      : yangkai
# Created On  : 2026-06-05
# ============================================================

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARCHIVE_DIR = ROOT / "docs/archive/2026-07-03-legacy-design-plans"
# 该脚本仅用于复现历史文档，禁止把输出重新标记为当前架构方案。
SRC = ARCHIVE_DIR / "Datalogue_设计开发方案_v2.0.docx"
OUT = ARCHIVE_DIR / "Datalogue_设计开发方案_v2.0_更新版_20260605.docx"


STATUS_ROWS = [
    [
        "数据集页面能力工作区",
        "数据集、表、字段、指标、维度、术语、验证、权限、版本统一治理。",
        "已完成：数据集列表 + 能力 Tab；数据表、字段标注、指标、维度、业务术语、分析蓝图、语义验证均有入口或实现。",
        "分析场景、权限、版本历史仍是入口/空态，需补齐业务模型。",
        "P0",
    ],
    [
        "字段标注与 AI 审核",
        "通过 AI 标注字段语义，作为语义层资产建设基础。",
        "已完成：按表分组、审核状态、候选筛选、AI 建议、置信度/原因/同义词/枚举值字段、状态回写。",
        "批量确认、批量转化、标注质量评估暂未实现。",
        "P1",
    ],
    [
        "字段快速转指标/维度",
        "从已审核字段快速沉淀指标和维度，避免重复建模。",
        "已完成：后端转换接口防重复，前端复用现有表单并同步已转化状态；名称优先使用 AI 建议。",
        "批量转换、冲突合并、推荐聚合质量提升待做。",
        "P1",
    ],
    [
        "业务术语治理",
        "统一业务别名、同义词、定义、资产绑定和冲突处理。",
        "MVP 已完成：术语 CRUD、资产绑定、AI/规则发现、冲突检测、变更日志；可绑定字段、指标、维度、分析蓝图。",
        "术语尚未接入问数运行时归一化链路，缺少审批流与影响分析。",
        "P0",
    ],
    [
        "分析蓝图",
        "沉淀复杂 SQL、固定分析路径、参数、输出列和触发问法，提升复杂问数确定性。",
        "MVP 已完成：模型、迁移、CRUD、异步 SQL 分析接口形态、四步向导、测试、发布、版本、回滚、使用统计/日志。",
        "AI SQL 分析和测试运行仍是规则/mock 兜底；尚未接入真实 LLM、真实只读执行、问数路由和权限控制。",
        "P0",
    ],
    [
        "语义验证",
        "在数据集内试问验证字段、指标、维度和语义配置效果。",
        "已保留：原试问验证入口与流程仍可访问。",
        "未把业务术语、分析蓝图纳入统一验证报告。",
        "P0",
    ],
    [
        "自然语言问数运行时",
        "LeadAgent + QueryGraph 执行 NL2DSL2SQL、分析、报告生成。",
        "部分已有基础：DSL/SQL 相关节点和校验设计存在。",
        "蓝图路由、术语归一化、知识库检索、反馈学习尚未形成完整闭环。",
        "P0",
    ],
    [
        "WeKnora 知识库",
        "把文档知识、指标口径、业务说明作为问数上下文。",
        "方案已有，部分配置入口/规划存在。",
        "运行时检索、数据集绑定、命中解释、知识回流未完成。",
        "P1",
    ],
    [
        "权限与版本历史",
        "控制数据集/资产可见性，并记录语义资产变更。",
        "入口已保留；蓝图自身已有版本快照和回滚。",
        "缺少数据集级权限模型、资产级权限、统一变更审计和版本对比。",
        "P0",
    ],
    [
        "分析场景",
        "把高频问题组织为面向业务的场景包。",
        "入口已保留。",
        "缺少场景模型、场景与蓝图/指标/术语的编排关系。",
        "P2",
    ],
]


TASK_ROWS = [
    [
        "T-001",
        "P0",
        "蓝图路由接入问数链路",
        "用户问题命中 active 分析蓝图后优先走蓝图参数抽取、测试执行和结果解释，而不是普通 Text2SQL。",
        "同一数据集内可按触发问法/关键词匹配蓝图；命中后生成参数草稿；执行日志记录 blueprint_id；未命中回退普通问数。",
    ],
    [
        "T-002",
        "P0",
        "真实蓝图 SQL 分析服务",
        "替换 analyze-sql mock，调用 LLM/规则解析 SQL 的参数、输出列、业务步骤、风险点和推荐触发问法。",
        "接口仍保持异步任务形态；支持失败状态；返回置信度和审核提示；后端测试覆盖成功/失败/低置信场景。",
    ],
    [
        "T-003",
        "P0",
        "真实蓝图测试执行",
        "把当前模拟测试替换为安全只读执行，支持参数校验、LIMIT、超时、脱敏和结果解释。",
        "未测试成功不能发布；执行失败返回诊断；使用日志包含耗时、行数、状态、错误摘要。",
    ],
    [
        "T-004",
        "P0",
        "业务术语运行时归一化",
        "在 QueryGraph 中加入 TermNormalizeNode，将业务术语、同义词和冲突信息注入 DSL/SQL 生成。",
        "同义词可影响字段/指标匹配；冲突术语进入澄清；语义验证可展示术语命中。",
    ],
    [
        "T-005",
        "P0",
        "数据集级权限模型",
        "补齐数据集、字段、指标、维度、术语、蓝图的可见性与操作权限。",
        "支持只读/编辑/发布角色；API 层校验；前端权限 Tab 可配置基础规则；测试覆盖越权访问。",
    ],
    [
        "T-006",
        "P0",
        "统一版本历史与变更审计",
        "把指标、维度、字段标注、业务术语、分析蓝图的创建/修改/发布汇总到数据集版本历史。",
        "版本历史 Tab 展示时间线、操作者、资产类型、变更摘要；支持蓝图版本跳转。",
    ],
    [
        "T-007",
        "P1",
        "字段批量审核与批量转化",
        "提升字段标注 Tab 的大表处理效率，支持按表/筛选结果批量确认、忽略、建指标、建维度。",
        "批量动作有二次确认；重复资产返回已存在并同步状态；前端无长列表卡顿。",
    ],
    [
        "T-008",
        "P1",
        "蓝图 AI 重分析与差异审核",
        "已发布蓝图编辑 SQL 或触发问法后，可重新分析并展示与当前版本的差异。",
        "差异视图覆盖参数、输出列、业务步骤、触发问法；用户可逐项接受/拒绝。",
    ],
    [
        "T-009",
        "P1",
        "WeKnora 运行时检索适配",
        "把知识库文档检索结果作为问数、术语解释和蓝图生成上下文。",
        "支持数据集绑定知识库；结果可追踪来源；低命中时降级不阻塞问数。",
    ],
    [
        "T-010",
        "P1",
        "LeadAgent/Skill 数据集级编排",
        "把指标、维度、术语、蓝图、知识库组织成可被 LeadAgent 调用的数据集 Skill。",
        "可生成数据集系统提示；复杂问题拆解时能选择普通问数、蓝图或知识检索子任务。",
    ],
    [
        "T-011",
        "P1",
        "术语冲突审核工作流",
        "在冲突检测基础上提供合并、保留、废弃和审批状态，避免术语库越用越乱。",
        "冲突项可处理；处理结果写变更日志；同义词冲突会影响问数澄清策略。",
    ],
    [
        "T-012",
        "P2",
        "分析场景模型与场景包",
        "把高频业务问题组织为场景，关联指标、维度、术语、蓝图和推荐问题。",
        "分析场景 Tab 可 CRUD；场景详情可查看关联资产和示例问法。",
    ],
    [
        "T-013",
        "P2",
        "API 发布与开放平台联动",
        "将稳定指标、蓝图和问数能力发布为可授权调用的 API。",
        "API 页面可选择语义资产发布；调用权限和审计记录可查询。",
    ],
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "D9E2EC", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill: str = "EAF2FF") -> None:
    table.autofit = False
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = RGBColor(23, 55, 94)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.add_run("• ").bold = True
    p.add_run(text)


def add_numbered(doc: Document, index: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.add_run(f"{index}. ").bold = True
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F8FF")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(20, 86, 143)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(43, 55, 72)
    doc.add_paragraph()


def add_status_table(doc: Document) -> None:
    widths = [1.15, 1.55, 2.1, 2.0, 0.55]
    table = doc.add_table(rows=1, cols=5)
    headers = ["模块", "v2.0 目标", "当前实现", "未完成差距", "优先级"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        table.rows[0].cells[i].width = Inches(widths[i])
    for row in STATUS_ROWS:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            cells[i].width = Inches(widths[i])
    style_table(table)


def add_task_table(doc: Document) -> None:
    widths = [0.58, 0.48, 1.45, 2.25, 2.25]
    table = doc.add_table(rows=1, cols=5)
    headers = ["编号", "优先级", "飞书任务标题", "目标", "验收标准"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        table.rows[0].cells[i].width = Inches(widths[i])
    for row in TASK_ROWS:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            cells[i].width = Inches(widths[i])
        priority = row[1]
        fill = {"P0": "FFE8E6", "P1": "FFF4D6", "P2": "EAF7EA"}.get(priority, "FFFFFF")
        set_cell_shading(cells[1], fill)
    style_table(table, header_fill="EEF2F7")


def append_update(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "附录 A：v2.0.1 更新 - 分析蓝图与实施差距", level=1)
    p = doc.add_paragraph()
    p.add_run("更新日期：").bold = True
    p.add_run("2026-06-05")
    p.add_run("    更新依据：").bold = True
    p.add_run("当前代码实现、数据集页面最新原型、分析蓝图 MVP 实施结果。")

    add_callout(
        doc,
        "更新结论",
        "分析蓝图已经从原 v2.0 方案中的泛化“分析配方/Skill”能力升级为数据集语义层的一等资产。"
        "它与指标、维度、业务术语一起构成复杂问数的确定性路径：普通问题走 NL2DSL2SQL，复杂固定路径问题优先匹配已发布蓝图。",
    )

    add_heading(doc, "A.1 产品规划调整", level=2)
    add_bullet(doc, "数据集页面从单点配置页调整为“数据集列表 + 能力 Tab 工作区”，所有语义治理能力在同一上下文内完成。")
    add_bullet(doc, "分析蓝图新增为语义资产：包含原始 SQL、参数、输出列、业务步骤、触发问法、测试记录、发布状态、版本快照和使用日志。")
    add_bullet(doc, "AI 字段标注不再止步于描述生成，而是进入“审核 → 快速创建指标/维度 → 状态回写”的资产沉淀闭环。")
    add_bullet(doc, "业务术语从静态词典升级为可绑定字段、指标、维度、分析蓝图的治理资产，并为后续问数归一化提供基础。")
    add_bullet(doc, "权限、版本历史、分析场景保留为数据集工作区入口，但当前仍需补齐后端模型和完整交互。")

    add_heading(doc, "A.2 分析蓝图架构定位", level=2)
    doc.add_paragraph(
        "分析蓝图用于承载复杂 SQL、固定口径分析、归因/漏斗/ cohort 等不适合每次临时生成 SQL 的场景。"
        "在运行时，系统应先判断用户问题是否命中已发布蓝图；若命中，则进入参数抽取、蓝图测试执行和结果解释；"
        "若未命中，再回退到普通 NL2DSL2SQL 链路。"
    )
    asset_table = doc.add_table(rows=1, cols=4)
    for i, header in enumerate(["语义资产", "作用", "当前状态", "下一步集成点"]):
        set_cell_text(asset_table.rows[0].cells[i], header, bold=True)
    asset_rows = [
        ["字段标注", "把物理字段映射为业务可理解字段。", "已实现审核与转化闭环。", "补批量审核和质量评分。"],
        ["指标/维度", "构建稳定 DSL/SQL 生成对象。", "CRUD 与快速创建已实现。", "接入术语归一化和版本历史。"],
        ["业务术语", "统一别名、同义词和定义。", "MVP 已实现。", "接入 QueryGraph 运行时。"],
        ["分析蓝图", "沉淀固定复杂分析路径。", "MVP 已实现。", "接入问数路由和真实执行。"],
        ["分析场景", "组织面向业务的高频问题包。", "入口保留。", "补模型与资产编排关系。"],
    ]
    for row in asset_rows:
        cells = asset_table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    style_table(asset_table)

    add_heading(doc, "A.3 当前实现状态与差距矩阵", level=2)
    doc.add_paragraph("下表按当前项目代码和已运行页面整理，作为后续排期和飞书任务创建依据。")
    add_status_table(doc)

    add_heading(doc, "A.4 建议创建的飞书任务", level=2)
    doc.add_paragraph(
        "优先级规则：P0 影响端到端闭环和准确率地基；P1 提升效率、质量和工程化能力；P2 面向企业化运营、开放平台和长期扩展。"
    )
    add_task_table(doc)

    add_heading(doc, "A.5 推荐实施顺序", level=2)
    steps = [
        "先完成 P0-1 至 P0-4：让分析蓝图和业务术语真正进入问数运行时，形成“复杂问题走蓝图、通用问题走 NL2SQL”的双路径。",
        "随后完成 P0-5 至 P0-6：补齐权限与统一版本历史，保证语义资产发布后可治理、可追踪、可回滚。",
        "再进入 P1：提升字段标注批量效率、蓝图重分析审核、WeKnora 检索和 LeadAgent/Skill 编排。",
        "最后推进 P2：分析场景、API 发布、运营监控、多租户等企业化能力。",
    ]
    for idx, text in enumerate(steps, 1):
        add_numbered(doc, idx, text)


def main() -> None:
    doc = Document(SRC)
    if len(doc.paragraphs) > 5:
        doc.paragraphs[5].text = "（v2.0.1 更新：新增分析蓝图一等语义资产 · 对齐当前实现 · 补充差距矩阵与飞书任务优先级）"

    for i, para in enumerate(doc.paragraphs):
        if i == 12 and "本 v2.0 方案" in para.text:
            para.text = (
                para.text
                + " 2026-06-05 追加 v2.0.1 更新：分析蓝图已纳入数据集语义治理核心资产，"
                "并补充当前实现状态、未完成能力和飞书任务优先级。"
            )
            break

    append_update(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
