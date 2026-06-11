# ============================================================
# File Name   : export_langfuse_docs_to_docx.py
# Description:
#   将 Langfuse 可观测能力 Markdown 文档导出为 Word 文档。
#
# Responsibilities:
#   - 解析需求设计与开发文档中的标题、段落、列表、表格和代码块。
#   - 生成适合团队评审和归档的 docx 版本。
#
# Author      : yangkai
# Created On  : 2026-06-10
# ============================================================

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCES = [
    ROOT / "docs" / "Langfuse可观测能力需求设计文档.md",
    ROOT / "docs" / "Langfuse可观测能力开发文档.md",
]
OUTPUT = ROOT / "docs" / "Langfuse可观测能力需求与开发文档.docx"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
ACCENT_BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
TABLE_HEADER_FILL = "F2F4F7"
CODE_FILL = "F7F9FB"


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin_name, value in CELL_MARGIN_DXA.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, column_widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(column_widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in column_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(column_widths[idx]))
            tc_w.set(qn("w:type"), "dxa")

    set_cell_margins(table)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_run_font(run, *, name: str = "Calibri", size: int | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)


def add_inline_runs(paragraph, text: str):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|<https?://[^>]+>)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas")
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif token.startswith("<"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run)
            run.font.color.rgb = ACCENT_BLUE
            run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT_BLUE, 16, 8),
        ("Heading 2", 13, ACCENT_BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("Heading 4", 11, DARK_BLUE, 6, 3),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    styles = doc.styles
    if "Code Block" not in [s.name for s in styles]:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(45, 55, 72)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.05


def configure_section(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("数语 Langfuse 可观测能力文档")
    set_run_font(run, size=9)
    run.font.color.rgb = MUTED


def add_cover(doc: Document):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(150)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Langfuse 可观测能力\n需求设计与开发文档")
    set_run_font(run, size=24)
    run.bold = True
    run.font.color.rgb = ACCENT_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run = subtitle.add_run("适用于数语 ChatBI + 多 Agent + 政企交付场景")
    set_run_font(run, size=12)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("项目", "数语（智能问数）"),
        ("文档类型", "需求设计文档 + 开发文档"),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("来源文件", "docs/Langfuse可观测能力需求设计文档.md；docs/Langfuse可观测能力开发文档.md"),
    ]
    for row, (key, value) in zip(meta.rows, rows):
        row.cells[0].text = key
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)
        shade_cell(row.cells[0], TABLE_HEADER_FILL)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_geometry(meta, [1800, PAGE_WIDTH_DXA - 1800])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_contents(doc: Document, sources: list[Path]):
    heading = doc.add_heading("文档结构", level=1)
    heading.runs[0].font.color.rgb = ACCENT_BLUE
    for idx, source in enumerate(sources, start=1):
        first_heading = next(
            (
                line.lstrip("#").strip()
                for line in source.read_text(encoding="utf-8").splitlines()
                if line.startswith("# ")
            ),
            source.stem,
        )
        p = doc.add_paragraph(style="List Number")
        add_inline_runs(p, f"{idx}. {first_heading}")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or not is_table_separator(lines[start + 1]):
        return None
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith("|") or not line.endswith("|"):
            break
        if not is_table_separator(line):
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        idx += 1
    return rows, idx


def column_widths(column_count: int) -> list[int]:
    if column_count <= 0:
        return [PAGE_WIDTH_DXA]
    if column_count == 2:
        return [2500, PAGE_WIDTH_DXA - 2500]
    if column_count == 3:
        return [1600, 2600, PAGE_WIDTH_DXA - 4200]
    if column_count == 4:
        return [1300, 1900, 2500, PAGE_WIDTH_DXA - 5700]
    base = PAGE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.allow_autofit = False

    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, text)
            if r_idx == 0:
                shade_cell(cell, TABLE_HEADER_FILL)
                for run in paragraph.runs:
                    run.bold = True
            if len(text) <= 8:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, column_widths(max_cols))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_code_block(doc: Document, code_lines: list[str]):
    for raw_line in code_lines:
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.08)
        shade_paragraph(p, CODE_FILL)
        run = p.add_run(raw_line if raw_line else " ")
        set_run_font(run, name="Consolas", size=9)


def add_markdown_file(doc: Document, path: Path, part_title: str):
    doc.add_heading(part_title, level=1)
    lines = path.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []

    while idx < len(lines):
        line = lines[idx]

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not line.strip():
            idx += 1
            continue

        table = parse_table(lines, idx)
        if table:
            rows, next_idx = table
            add_markdown_table(doc, rows)
            idx = next_idx
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)) + 1, 4)
            text = heading_match.group(2).strip()
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                set_run_font(run)
            idx += 1
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet_match.group(1).strip())
            idx += 1
            continue

        ordered_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if ordered_match:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, ordered_match.group(1).strip())
            idx += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line.strip())
        idx += 1


def build_docx():
    doc = Document()
    configure_section(doc)
    configure_styles(doc)
    add_footer(doc)
    add_cover(doc)
    add_contents(doc, SOURCES)
    add_markdown_file(doc, SOURCES[0], "第一部分：需求设计文档")
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_markdown_file(doc, SOURCES[1], "第二部分：开发文档")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
